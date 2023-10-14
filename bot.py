import discord
import pymongo
import datetime
import re
import urllib.request
import requests
from bs4 import BeautifulSoup
from discord.ext import commands, tasks
from discord.ext.commands import Bot
import os
import asyncio
import re
import docx
from pymongo import MongoClient
import tweepy  


intents = discord.Intents.all()
intents.members = True
intents.typing = True
intents.presences = True
bot = commands.Bot(command_prefix='!', intents=intents)

client = pymongo.MongoClient("mongodb://localhost:27017/")
db = client["botdatabase"]

registrations = db["registrations"]
eligibility = db["eligibility"]

consumer_key = 'your_consumer_key'
consumer_secret = 'your_consumer_secret'
access_token = 'your_access_token'
access_token_secret = 'your_access_token_secret'

auth = tweepy.OAuthHandler(consumer_key, consumer_secret)
auth.set_access_token(access_token, access_token_secret)
api = tweepy.API(auth)

faq_dict = {
    "Q1": "What is this bot for?",
    "A1": "This bot helps you register for events, check your eligibility, participate in challenges, and more!",
    "Q2": "How do I register for an event?",
    "A2": "You can register for an event by using the `!register` command followed by the post link.",
    "Q3": "How can I check my eligibility?",
    "A3": "You can check your eligibility by using the `!check_eligibility` command.",
    "Q4": "Can I participate in multiple events simultaneously?",
    "A4": "Yes, you can participate in multiple events at the same time.",
    "Q5": "What is the minimum post requirement for eligibility?",
    "A5": "The minimum post requirement varies for each event. Check the event details for the specific requirement.",
    "Q6": "How are tokens distributed in events?",
    "A6": "Tokens are distributed to eligible participants by administrators using the `!distribute_tokens` command.",
    
    }


@bot.command()
async def faq(ctx, question: str):
    if question in faq_dict:
        answer_key = f"A{question[1:]}" 
        if answer_key in faq_dict:
            answer = faq_dict[answer_key]
            await ctx.send(f"**{question}:** {answer}")
        else:
            await ctx.send("The answer for this question is not available.")
    else:
        await ctx.send("This question is not in the FAQ. Please check the available FAQ questions.")

@bot.command()
async def listfaq(ctx):
    faq_list = "\n".join([f"**{key}:** {value}" for key, value in faq_dict.items()])
    await ctx.send(f"Available FAQ questions and answers:\n{faq_list}")
@bot.command()
async def greeting(ctx):
    await ctx.send('Hello! my name is event-verifier have a good day!'  )

@bot.command()
async def about(ctx):
    await ctx.send('I am a Discord bot that verifies the post for different challenges and maintains your streak! and help you complete the challenge')

@bot.event
async def on_ready():
    print(f'Logged in as {bot.user.name} ({bot.user.id})')

def get_tweet_creation_date(tweet_url):
    tweet_id = tweet_url.split('/')[-1]
    tweet = api.get_status(tweet_id)
    return tweet.created_at

@bot.command()
async def post(ctx, post_link):
    user_id = ctx.author.id
    username = ctx.author.name
    timestamp = datetime.datetime.now()
    last_registration = registrations.find_one({"user_id": user_id}, sort=[("timestamp", pymongo.DESCENDING)])
    
    if last_registration:
        last_timestamp = last_registration["timestamp"]
        if (timestamp - last_timestamp).days == 0:
            await ctx.send("You have already registered for today.")
            return
    
    if not (post_link.startswith('https://twitter.com') or post_link.startswith('https://linkedin.com')):
        await ctx.send("Invalid format. Please provide a valid Twitter or LinkedIn post link.")
        return

    registration_data = {
        "user_id": user_id,
        "username": username,
        "post_link": post_link,
        "timestamp": timestamp
    }
    registrations.insert_one(registration_data)
    await ctx.send(f"Registration complete, {ctx.author.mention}!")
    tweet_creation_date = get_tweet_creation_date(post_link)
    if tweet_creation_date is not None:
        days_old = (timestamp - tweet_creation_date).days
        if days_old > 1: 
            await ctx.send("This post is too old to be eligible.")
            return

    

@bot.command()
async def daily_check(ctx):
    yesterday = datetime.datetime.now() - datetime.timedelta(days=1)
    registered_users_for_today = list(registrations.find({"timestamp": {"$gte": yesterday}}))
    
    for registration in registered_users_for_today:
        user_id = registration["user_id"]
        user_registration = registrations.find_one({"user_id": user_id, "timestamp": {"$gte": yesterday}})
        
        if not user_registration:
            eligibility_data = {
                "user_id": user_id,
                "eligible": False
            }
        else:
            post_link = user_registration["post_link"]
            tweet_creation_date = get_tweet_creation_date(post_link)
            
            if tweet_creation_date is not None:
                days_old = (datetime.datetime.now() - tweet_creation_date).days
                if days_old <= 7:  # Consider posts within the last 7 days
                    if all(keyword in post_link for keyword in ["#Opensourcesep", "#ScalerDiscord", "@Scaler_Official"]):
                        eligibility_data = {
                            "user_id": user_id,
                            "eligible": True
                        }
                    else:
                        eligibility_data = {
                            "user_id": user_id,
                            "eligible": False
                        }
                else:
                    eligibility_data = {
                        "user_id": user_id,
                        "eligible": False  # Mark old posts as ineligible
                    }
            
            eligibility.update_one({"user_id": user_id}, {"$set": {"eligible": eligibility_data["eligible"]}}, upsert=True)

    print("executed")

@bot.command()
async def check_eligibility(ctx):
    user_id = ctx.author.id
    result = eligibility.find_one({"user_id": user_id})
    if result:
        eligible = result["eligible"]
        if eligible:
            await ctx.send("You are eligible for rewards.")
        else:
            await ctx.send("You are ineligible for rewards.")
    else:
        await ctx.send("You have not registered for the challenge.")

@bot.event
async def on_error(event, *args, **kwargs):
    pass

@bot.event
async def on_command_error(ctx, error):
    if isinstance(error, discord.ext.commands.CommandNotFound):
        await ctx.send("Command not found. Use !help for a list of available commands.")
    else:
        await ctx.send(f"An error occurred: {str(error)}")
reminder_flags={}        
@bot.command()
async def user_reminder(ctx):
    user_id = ctx.author.id

    if user_id not in reminder_flags or not reminder_flags[user_id]:
        reminder_flags[user_id] = True
        await ctx.send("You will receive daily reminders to post.")
        await post_reminder(ctx, user_id)
    else:
        await ctx.send("You are already scheduled to receive reminders.")

@bot.command()
async def decline_reminder(ctx):
    user_id = ctx.author.id
    if user_id in reminder_flags and reminder_flags[user_id]:
        reminder_flags[user_id] = False
        await ctx.send("You will no longer receive daily reminders.")
    else:
        await ctx.send("You don't have any active reminders.")

async def post_reminder(ctx, user_id):
    while reminder_flags.get(user_id, False):
        await ctx.send("Don't forget to post today!")
        await asyncio.sleep(86400)  
        
@bot.command()
async def distribute_tokens(ctx, amount_per_user: int):
    if any(role.name == "Administrator" for role in ctx.author.roles):
        eligible_users = list(db.eligibility.find({"eligibility": True}))
        total_users = len(eligible_users)
        if total_users > 0:
            total_tokens = amount_per_user * total_users  # Total tokens to distribute
            failed_users = []

            for user_data in eligible_users:
                user_id = user_data["user_id"]
                try:
                    if user_id in user_tokens:
                        user_tokens[user_id] += amount_per_user  # Add tokens to the user's balance
                    else:
                        user_tokens[user_id] = amount_per_user  # Create a new balance for the user
                except Exception as e:
                    error_message = f"Error distributing tokens to user {user_id}: {str(e)}"
                    print(error_message)
                    await ctx.send(error_message)  # Send an error message to the user
                    failed_users.append(user_id)

            if failed_users:
                await ctx.send(f"{total_tokens} tokens distributed to eligible participants, but there were errors for some users.")
            else:
                await ctx.send(f"{total_tokens} tokens distributed to eligible participants successfully.")
        else:
            await ctx.send("There are no eligible participants based on the 'eligibility' table.")
    else:
        await ctx.send("You do not have permission to distribute tokens.")


@bot.command()
async def download_eligible(ctx):
    if any(role.name == "Administrator" for role in ctx.author.roles):
        cursor = eligibility.find({}, {"_id": 0, "user_id": 1, "eligible": 1})
        eligible_users_list = list(cursor)
        if not eligible_users_list:
            await ctx.send("There are no eligible participants in the database.")
            return
        document = docx.Document()
        document.add_heading("Eligible Participants", 0)
        table = document.add_table(rows=1, cols=2)
        table.autofit = True
        table.style = "Table Grid"
        table.allow_autofit = False
        table.cell(0, 0).text = "User ID"
        table.cell(0, 1).text = "Eligibility"
        for i, user_data in enumerate(eligible_users_list, 1):
            user_id = str(user_data["user_id"])
            eligibility = "Eligible" if user_data["eligible"] else "Ineligible"
            row_cells = table.add_row().cells
            row_cells[0].text = user_id
            row_cells[1].text = eligibility
        document.save("eligible_participants.docx")
        with open("eligible_participants.docx", "rb") as file:
            await ctx.send("Here is the list of eligible participants:", file=discord.File(file, "eligible_participants.docx"))
        os.remove("eligible_participants.docx")

bot.run("MTE2MDEzNzg4MTQyNDI0NDc2Nw.GiZ8x0.hcvT4PawJqYnmJpB37wbIPrwL0FzUJ6QXJh3Wg")

